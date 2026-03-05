#!/usr/bin/env python3
"""
Batch Document Processing Module
===============================
Advanced batch processor for PDFs, XML, Office documents, and various file formats.
"""

import os
import json
import time
import xml.etree.ElementTree as ET
import xml.dom.minidom as minidom
from pathlib import Path
from typing import List, Dict, Tuple, Optional, Union
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
import warnings
import re
import zipfile
import hashlib
import mimetypes
warnings.filterwarnings('ignore')

try:
    import pandas as pd
    import chardet
    from tqdm import tqdm
    CORE_AVAILABLE = True
except ImportError:
    CORE_AVAILABLE = False
    print("⚠️ Core libraries not available. Install: pip install pandas chardet tqdm")

try:
    from bs4 import BeautifulSoup
    HTML_PROCESSING_AVAILABLE = True
except ImportError:
    HTML_PROCESSING_AVAILABLE = False
    print("⚠️ HTML processing not available. Install: pip install beautifulsoup4")

try:
    import PyPDF2
    import pdfplumber
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    print("⚠️ PDF processing not available. Install: pip install PyPDF2 pdfplumber")

try:
    from docx import Document
    DOCX_PROCESSING_AVAILABLE = True
except ImportError:
    DOCX_PROCESSING_AVAILABLE = False
    print("⚠️ DOCX processing not available. Install: pip install python-docx")

try:
    import openpyxl
    import xlrd
    EXCEL_PROCESSING_AVAILABLE = True
except ImportError:
    EXCEL_PROCESSING_AVAILABLE = False
    print("⚠️ Excel processing not available. Install: pip install openpyxl xlrd")

try:
    from lxml import etree
    ADVANCED_XML_AVAILABLE = True
except ImportError:
    ADVANCED_XML_AVAILABLE = False
    print("⚠️ Advanced XML processing not available. Install: pip install lxml")

class BatchDocumentProcessor:
    """Advanced batch processor for documents and files"""
    
    def __init__(self, output_dir="document_output"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        # Supported file formats
        self.supported_formats = {
            'pdf': ['.pdf'],
            'office': ['.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx'],
            'xml': ['.xml', '.xsd', '.xsl', '.xslt'],
            'text': ['.txt', '.rtf', '.csv'],
            'web': ['.html', '.htm', '.xhtml'],
            'archive': ['.zip', '.rar', '.7z'],
            'email': ['.eml', '.msg'],
            'open_office': ['.odt', '.ods', '.odp']
        }
        
        # Processing statistics
        self.stats = {
            'processed': 0,
            'errors': 0,
            'skipped': 0,
            'total_size': 0,
            'processing_time': 0,
            'by_format': {}
        }
        
        # Initialize MIME type detection
        mimetypes.init()
    
    def detect_file_type(self, file_path: Path) -> Dict[str, str]:
        """Detect file type and encoding"""
        try:
            # Basic file info
            info = {
                'extension': file_path.suffix.lower(),
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime)
            }
            
            # MIME type detection
            mime_type, encoding = mimetypes.guess_type(str(file_path))
            info['mime_type'] = mime_type
            info['encoding'] = encoding
            
            # Determine category
            category = 'unknown'
            for cat, extensions in self.supported_formats.items():
                if info['extension'] in extensions:
                    category = cat
                    break
            info['category'] = category
            
            # For text files, detect encoding
            if category in ['text', 'xml', 'web'] and CORE_AVAILABLE:
                try:
                    with open(file_path, 'rb') as f:
                        raw_data = f.read(10000)  # Read first 10KB
                        detected = chardet.detect(raw_data)
                        info['detected_encoding'] = detected.get('encoding', 'utf-8')
                        info['encoding_confidence'] = detected.get('confidence', 0)
                except:
                    info['detected_encoding'] = 'utf-8'
                    info['encoding_confidence'] = 0
            
            return info
        
        except Exception as e:
            return {'error': str(e), 'path': str(file_path)}
    
    def extract_pdf_content(self, pdf_path: Path) -> Dict:
        """Extract comprehensive content from PDF"""
        if not PDF_PROCESSING_AVAILABLE:
            return {'error': 'PDF processing libraries not available'}
            
        try:
            content = {
                'text': '',
                'metadata': {},
                'pages': 0,
                'images': [],
                'tables': [],
                'bookmarks': [],
                'annotations': []
            }
            
            # Extract with pdfplumber (better for text and tables)
            with pdfplumber.open(pdf_path) as pdf:
                content['pages'] = len(pdf.pages)
                content['metadata'] = pdf.metadata or {}
                
                # Extract text and tables from each page
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text() or ''
                    content['text'] += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    
                    # Extract tables
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        content['tables'].append({
                            'page': page_num + 1,
                            'table_num': table_num + 1,
                            'data': table,
                            'rows': len(table),
                            'cols': len(table[0]) if table else 0
                        })
            
            # Extract additional info with PyPDF2
            try:
                with open(pdf_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    
                    # Extract bookmarks
                    if reader.outline:
                        content['bookmarks'] = self._extract_bookmarks(reader.outline)
                    
                    # Extract annotations
                    for page_num, page in enumerate(reader.pages):
                        if '/Annots' in page:
                            annotations = page['/Annots']
                            for annot in annotations:
                                annot_obj = annot.get_object()
                                if '/Contents' in annot_obj:
                                    content['annotations'].append({
                                        'page': page_num + 1,
                                        'type': str(annot_obj.get('/Subtype', '')),
                                        'content': str(annot_obj.get('/Contents', ''))
                                    })
            except:
                pass  # Continue if PyPDF2 extraction fails
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(pdf_path)}
    
    def _extract_bookmarks(self, outline, level=0):
        """Recursively extract PDF bookmarks"""
        bookmarks = []
        for item in outline:
            if isinstance(item, list):
                bookmarks.extend(self._extract_bookmarks(item, level + 1))
            else:
                bookmarks.append({
                    'title': item.title,
                    'level': level,
                    'page': item.page.idnum if hasattr(item, 'page') else None
                })
        return bookmarks
    
    def extract_docx_content(self, docx_path: Path) -> Dict:
        """Extract content from DOCX files"""
        if not DOCX_PROCESSING_AVAILABLE:
            return {'error': 'DOCX processing libraries not available'}
            
        try:
            doc = Document(docx_path)
            
            content = {
                'text': '',
                'paragraphs': [],
                'tables': [],
                'images': [],
                'headers': [],
                'footers': [],
                'metadata': {}
            }
            
            # Extract core properties
            core_props = doc.core_properties
            content['metadata'] = {
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'created': core_props.created,
                'modified': core_props.modified,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision
            }
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append({
                        'text': para.text,
                        'style': para.style.name if para.style else None
                    })
                    content['text'] += para.text + '\n'
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                content['tables'].append({
                    'table_num': table_num + 1,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(docx_path)}
    
    def extract_xml_content(self, xml_path: Path) -> Dict:
        """Extract and validate XML content"""
        try:
            content = {
                'text': '',
                'structure': {},
                'elements': [],
                'attributes': {},
                'namespaces': {},
                'validation': {'valid': False, 'errors': []}
            }
            
            # Read and parse XML
            with open(xml_path, 'r', encoding='utf-8') as f:
                xml_content = f.read()
                content['text'] = xml_content
            
            # Parse with ElementTree
            try:
                root = ET.fromstring(xml_content)
                content['validation']['valid'] = True
                
                # Extract structure
                content['structure'] = self._analyze_xml_structure(root)
                
                # Extract all elements
                for elem in root.iter():
                    content['elements'].append({
                        'tag': elem.tag,
                        'text': elem.text.strip() if elem.text else '',
                        'attributes': dict(elem.attrib),
                        'children': len(list(elem))
                    })
                
                # Extract namespaces
                content['namespaces'] = self._extract_namespaces(xml_content)
                
            except ET.ParseError as e:
                content['validation']['valid'] = False
                content['validation']['errors'].append(str(e))
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(xml_path)}
    
    def _analyze_xml_structure(self, element, level=0):
        """Analyze XML structure recursively"""
        structure = {
            'tag': element.tag,
            'level': level,
            'attributes': len(element.attrib),
            'has_text': bool(element.text and element.text.strip()),
            'children': []
        }
        
        for child in element:
            structure['children'].append(self._analyze_xml_structure(child, level + 1))
        
        return structure
    
    def _extract_namespaces(self, xml_content):
        """Extract XML namespaces"""
        namespaces = {}
        
        # Find namespace declarations
        ns_pattern = r'xmlns(?::(\w+))?=["\'](.*?)["\']'
        matches = re.findall(ns_pattern, xml_content)
        
        for prefix, uri in matches:
            if prefix:
                namespaces[prefix] = uri
            else:
                namespaces['default'] = uri
        
        return namespaces
    
    def extract_excel_content(self, excel_path: Path) -> Dict:
        """Extract content from Excel files"""
        if not EXCEL_PROCESSING_AVAILABLE or not CORE_AVAILABLE:
            return {'error': 'Excel processing libraries not available'}
            
        try:
            content = {
                'sheets': [],
                'metadata': {},
                'total_rows': 0,
                'total_cols': 0
            }
            
            # Read with pandas for easy processing
            excel_file = pd.ExcelFile(excel_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                
                sheet_info = {
                    'name': sheet_name,
                    'rows': len(df),
                    'cols': len(df.columns),
                    'columns': df.columns.tolist(),
                    'data_types': df.dtypes.to_dict(),
                    'sample_data': df.head().to_dict('records') if not df.empty else []
                }
                
                content['sheets'].append(sheet_info)
                content['total_rows'] += len(df)
                content['total_cols'] = max(content['total_cols'], len(df.columns))
            
            # Extract metadata using openpyxl for XLSX files
            if excel_path.suffix.lower() == '.xlsx':
                try:
                    wb = openpyxl.load_workbook(excel_path)
                    content['metadata'] = {
                        'title': wb.properties.title,
                        'creator': wb.properties.creator,
                        'created': wb.properties.created,
                        'modified': wb.properties.modified,
                        'last_modified_by': wb.properties.lastModifiedBy
                    }
                except:
                    pass
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(excel_path)}
    
    def extract_archive_content(self, archive_path: Path) -> Dict:
        """Extract content information from archives"""
        try:
            content = {
                'files': [],
                'total_files': 0,
                'total_size': 0,
                'compression_ratio': 0,
                'file_types': {}
            }
            
            if archive_path.suffix.lower() == '.zip':
                with zipfile.ZipFile(archive_path, 'r') as zip_file:
                    for file_info in zip_file.filelist:
                        file_data = {
                            'filename': file_info.filename,
                            'file_size': file_info.file_size,
                            'compress_size': file_info.compress_size,
                            'date_time': datetime(*file_info.date_time),
                            'is_dir': file_info.is_dir()
                        }
                        
                        if not file_info.is_dir():
                            content['files'].append(file_data)
                            content['total_files'] += 1
                            content['total_size'] += file_info.file_size
                            
                            # Track file types
                            ext = Path(file_info.filename).suffix.lower()
                            content['file_types'][ext] = content['file_types'].get(ext, 0) + 1
                    
                    # Calculate compression ratio
                    if content['total_size'] > 0:
                        compressed_size = sum(f['compress_size'] for f in content['files'])
                        content['compression_ratio'] = compressed_size / content['total_size']
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(archive_path)}
    
    def extract_text_content(self, text_path: Path, file_info: Dict) -> Dict:
        """Extract content from text files"""
        try:
            encoding = file_info.get('detected_encoding', 'utf-8')
            with open(text_path, 'r', encoding=encoding) as f:
                content_text = f.read()
            
            content = {
                'text': content_text,
                'lines': len(content_text.splitlines()),
                'words': len(content_text.split()),
                'characters': len(content_text),
                'encoding_used': encoding
            }
            
            # Additional analysis for CSV files
            if text_path.suffix.lower() == '.csv' and CORE_AVAILABLE:
                try:
                    df = pd.read_csv(text_path)
                    content['csv_info'] = {
                        'rows': len(df),
                        'columns': len(df.columns),
                        'column_names': df.columns.tolist(),
                        'data_types': df.dtypes.to_dict()
                    }
                except:
                    pass
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(text_path)}
    
    def extract_html_content(self, html_path: Path, file_info: Dict) -> Dict:
        """Extract content from HTML files"""
        if not HTML_PROCESSING_AVAILABLE:
            return self.extract_text_content(html_path, file_info)
            
        try:
            encoding = file_info.get('detected_encoding', 'utf-8')
            with open(html_path, 'r', encoding=encoding) as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            content = {
                'text': soup.get_text(),
                'title': soup.title.string if soup.title else '',
                'meta_tags': [],
                'links': [],
                'images': [],
                'scripts': [],
                'stylesheets': []
            }
            
            # Extract meta tags
            for meta in soup.find_all('meta'):
                content['meta_tags'].append(dict(meta.attrs))
            
            # Extract links
            for link in soup.find_all('a', href=True):
                content['links'].append({
                    'href': link['href'],
                    'text': link.get_text().strip()
                })
            
            # Extract images
            for img in soup.find_all('img', src=True):
                content['images'].append({
                    'src': img['src'],
                    'alt': img.get('alt', '')
                })
            
            # Extract scripts
            for script in soup.find_all('script', src=True):
                content['scripts'].append(script['src'])
            
            # Extract stylesheets
            for link in soup.find_all('link', rel='stylesheet'):
                if link.get('href'):
                    content['stylesheets'].append(link['href'])
            
            return content
        
        except Exception as e:
            return {'error': str(e), 'path': str(html_path)}
    
    def process_single_document(self, doc_path: Path) -> Dict:
        """Process a single document"""
        try:
            start_time = time.time()
            
            # Detect file type
            file_info = self.detect_file_type(doc_path)
            
            if 'error' in file_info:
                return file_info
            
            # Initialize result
            result = {
                'success': True,
                'source': str(doc_path),
                'file_info': file_info,
                'content': {},
                'processing_time': 0
            }
            
            # Process based on category
            category = file_info['category']
            
            if category == 'pdf':
                result['content'] = self.extract_pdf_content(doc_path)
            elif category == 'office' and doc_path.suffix.lower() in ['.doc', '.docx']:
                result['content'] = self.extract_docx_content(doc_path)
            elif category == 'office' and doc_path.suffix.lower() in ['.xls', '.xlsx']:
                result['content'] = self.extract_excel_content(doc_path)
            elif category == 'xml':
                result['content'] = self.extract_xml_content(doc_path)
            elif category == 'archive':
                result['content'] = self.extract_archive_content(doc_path)
            elif category == 'text':
                result['content'] = self.extract_text_content(doc_path, file_info)
            elif category == 'web':
                result['content'] = self.extract_html_content(doc_path, file_info)
            else:
                result['content'] = {'message': f'Processing for {category} not implemented'}
            
            # Calculate processing time
            result['processing_time'] = time.time() - start_time
            
            # Update statistics
            self.stats['processed'] += 1
            self.stats['total_size'] += file_info['size']
            self.stats['processing_time'] += result['processing_time']
            
            if category not in self.stats['by_format']:
                self.stats['by_format'][category] = 0
            self.stats['by_format'][category] += 1
            
            return result
        
        except Exception as e:
            self.stats['errors'] += 1
            return {
                'success': False,
                'source': str(doc_path),
                'error': str(e)
            }
    
    def scan_directory(self, directory: Path, recursive: bool = True) -> List[Path]:
        """Scan directory for supported documents"""
        document_files = []
        
        if recursive:
            pattern = "**/*"
        else:
            pattern = "*"
        
        all_extensions = []
        for ext_list in self.supported_formats.values():
            all_extensions.extend(ext_list)
        
        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in all_extensions:
                document_files.append(file_path)
        
        return sorted(document_files)
    
    def batch_process_directory(self, input_dir: Path, recursive: bool = True, 
                              max_workers: int = 4) -> List[Dict]:
        """Batch process all documents in directory"""
        # Scan for documents
        document_files = self.scan_directory(input_dir, recursive)
        
        if not document_files:
            print(f"No supported documents found in {input_dir}")
            return []
        
        print(f"Found {len(document_files)} documents to process")
        
        # Reset statistics
        self.stats = {key: 0 if isinstance(self.stats[key], (int, float)) else {} 
                     for key in self.stats}
        
        results = []
        
        # Process with progress bar
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = []
            for doc_path in document_files:
                future = executor.submit(self.process_single_document, doc_path)
                futures.append(future)
            
            # Collect results with progress bar
            if CORE_AVAILABLE:
                try:
                    for future in tqdm(futures, desc="Processing documents"):
                        result = future.result()
                        results.append(result)
                except ImportError:
                    # Fallback without progress bar
                    for future in futures:
                        result = future.result()
                        results.append(result)
            else:
                for future in futures:
                    result = future.result()
                    results.append(result)
        
        return results
    
    def generate_processing_report(self, results: List[Dict]) -> Dict:
        """Generate comprehensive processing report"""
        successful = [r for r in results if r.get('success', False)]
        failed = [r for r in results if not r.get('success', False)]
        
        # Analyze by format
        format_analysis = {}
        for result in successful:
            category = result.get('file_info', {}).get('category', 'unknown')
            if category not in format_analysis:
                format_analysis[category] = {
                    'count': 0,
                    'total_size': 0,
                    'avg_processing_time': 0,
                    'files': []
                }
            
            format_analysis[category]['count'] += 1
            format_analysis[category]['total_size'] += result.get('file_info', {}).get('size', 0)
            format_analysis[category]['avg_processing_time'] += result.get('processing_time', 0)
            format_analysis[category]['files'].append(result['source'])
        
        # Calculate averages
        for category in format_analysis:
            count = format_analysis[category]['count']
            if count > 0:
                format_analysis[category]['avg_processing_time'] /= count
        
        report = {
            'summary': {
                'total_processed': len(results),
                'successful': len(successful),
                'failed': len(failed),
                'success_rate': len(successful) / len(results) * 100 if results else 0,
                'total_size_mb': self.stats['total_size'] / (1024 * 1024),
                'total_processing_time': self.stats['processing_time'],
                'avg_processing_time': self.stats['processing_time'] / len(successful) if successful else 0
            },
            'by_format': format_analysis,
            'successful_files': [r['source'] for r in successful],
            'failed_files': [(r['source'], r.get('error', 'Unknown error')) for r in failed],
            'statistics': self.stats
        }
        
        # Save report to file
        report_path = self.output_dir / "document_processing_report.json"
        with open(report_path, 'w') as f:
            json.dump(report, f, indent=2, default=str)
        
        return report

def create_sample_documents(output_dir="sample_documents"):
    """Create sample documents for testing"""
    sample_dir = Path(output_dir)
    sample_dir.mkdir(exist_ok=True)
    
    created_files = []
    
    # Create sample text file
    text_content = """Sample Text Document
====================

This is a sample text document for testing the batch document processor.

Features:
- Text extraction
- Line counting
- Word analysis
- Character counting

This document contains multiple paragraphs and demonstrates
how the processor handles plain text files.
"""
    
    text_file = sample_dir / "sample.txt"
    with open(text_file, 'w', encoding='utf-8') as f:
        f.write(text_content)
    created_files.append(text_file)
    
    # Create sample XML file
    xml_content = """<?xml version="1.0" encoding="UTF-8"?>
<document xmlns:meta="http://example.com/metadata">
    <meta:info>
        <meta:title>Sample XML Document</meta:title>
        <meta:author>Document Processor</meta:author>
        <meta:created>2024-01-01</meta:created>
    </meta:info>
    <content>
        <section id="introduction">
            <title>Introduction</title>
            <paragraph>This is a sample XML document for testing.</paragraph>
            <list>
                <item>XML parsing</item>
                <item>Structure analysis</item>
                <item>Namespace extraction</item>
            </list>
        </section>
        <section id="features">
            <title>Features</title>
            <paragraph>The processor can extract various XML elements.</paragraph>
        </section>
    </content>
</document>
"""
    
    xml_file = sample_dir / "sample.xml"
    with open(xml_file, 'w', encoding='utf-8') as f:
        f.write(xml_content)
    created_files.append(xml_file)
    
    # Create sample HTML file
    html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="author" content="Document Processor">
    <title>Sample HTML Document</title>
    <link rel="stylesheet" href="styles.css">
</head>
<body>
    <h1>Sample HTML Document</h1>
    <p>This is a sample HTML document for testing the batch processor.</p>
    
    <h2>Features</h2>
    <ul>
        <li>HTML parsing</li>
        <li>Metadata extraction</li>
        <li>Link analysis</li>
        <li>Image detection</li>
    </ul>
    
    <p>Visit <a href="https://example.com">Example</a> for more information.</p>
    
    <img src="sample.jpg" alt="Sample Image">
    
    <script src="script.js"></script>
</body>
</html>
"""
    
    html_file = sample_dir / "sample.html"
    with open(html_file, 'w', encoding='utf-8') as f:
        f.write(html_content)
    created_files.append(html_file)
    
    # Create sample CSV file
    if CORE_AVAILABLE:
        csv_data = {
            'Name': ['Alice', 'Bob', 'Charlie', 'Diana'],
            'Age': [25, 30, 35, 28],
            'City': ['New York', 'London', 'Paris', 'Tokyo'],
            'Score': [85.5, 92.3, 78.9, 88.7]
        }
        
        df = pd.DataFrame(csv_data)
        csv_file = sample_dir / "sample.csv"
        df.to_csv(csv_file, index=False)
        created_files.append(csv_file)
    
    return created_files

if __name__ == "__main__":
    # Demo usage
    print("🚀 Document Processor Demo")
    print("=" * 30)
    
    # Create sample documents
    sample_files = create_sample_documents()
    print(f"✅ Created {len(sample_files)} sample documents")
    
    # Initialize processor
    processor = BatchDocumentProcessor()
    
    # Process sample documents
    results = processor.batch_process_directory(Path("sample_documents"))
    
    # Generate report
    report = processor.generate_processing_report(results)
    
    print(f"\n📊 Processing completed:")
    print(f"   • Processed: {report['summary']['successful']}/{report['summary']['total_processed']}")
    print(f"   • Success rate: {report['summary']['success_rate']:.1f}%")
    print(f"   • Total size: {report['summary']['total_size_mb']:.2f} MB")
    print(f"   • Output directory: {processor.output_dir}")
    
    # Show format breakdown
    print(f"\n📂 By format:")
    for format_name, info in report['by_format'].items():
        print(f"   • {format_name}: {info['count']} files")
