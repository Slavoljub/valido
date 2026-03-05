#!/usr/bin/env python3
"""
Content Management System for ValidoAI
=====================================

Comprehensive file and content management with upload, processing, and serving capabilities.
Supports various content types: documents, images, videos, audio, and structured data.

Features:
- File upload and validation
- Content processing and analysis
- Metadata extraction and storage
- File serving and streaming
- Search and organization
- Access control and permissions
- Content versioning
- CDN integration support
"""

import os
import json
import logging
import hashlib
import mimetypes
import tempfile
import shutil
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, BinaryIO
import uuid
import magic  # for better file type detection
from werkzeug.utils import secure_filename
from PIL import Image
import PyPDF2
import docx
from pathlib import Path

logger = logging.getLogger(__name__)

class ContentManager:
    """Advanced content management system"""

    def __init__(self, upload_folder: str = "uploads", max_file_size: int = 50 * 1024 * 1024):
        self.upload_folder = Path(upload_folder)
        self.max_file_size = max_file_size
        self.temp_folder = Path("temp/content")
        self.cache_folder = Path("cache/content")

        # Create directories
        self.upload_folder.mkdir(parents=True, exist_ok=True)
        self.temp_folder.mkdir(parents=True, exist_ok=True)
        self.cache_folder.mkdir(parents=True, exist_ok=True)

        # Initialize content database (in production, use actual database)
        self.content_db = {}
        self._load_content_database()

        logger.info(f"✅ Content Manager initialized with upload folder: {self.upload_folder}")

    def _load_content_database(self):
        """Load content database from disk"""
        db_path = self.upload_folder / "content_db.json"
        if db_path.exists():
            try:
                with open(db_path, 'r', encoding='utf-8') as f:
                    self.content_db = json.load(f)
                logger.info(f"✅ Loaded content database with {len(self.content_db)} items")
            except Exception as e:
                logger.error(f"❌ Error loading content database: {e}")
                self.content_db = {}

    def _save_content_database(self):
        """Save content database to disk"""
        db_path = self.upload_folder / "content_db.json"
        try:
            with open(db_path, 'w', encoding='utf-8') as f:
                json.dump(self.content_db, f, indent=2, ensure_ascii=False)
        except Exception as e:
            logger.error(f"❌ Error saving content database: {e}")

    def validate_file(self, file, allowed_extensions: List[str] = None) -> Dict[str, Any]:
        """Validate uploaded file"""
        try:
            # Check if file exists
            if not hasattr(file, 'filename') or not file.filename:
                return {'valid': False, 'error': 'No file provided'}

            # Get secure filename
            filename = secure_filename(file.filename)
            if not filename:
                return {'valid': False, 'error': 'Invalid filename'}

            # Check file size
            file.seek(0, os.SEEK_END)
            file_size = file.tell()
            file.seek(0)

            if file_size > self.max_file_size:
                return {
                    'valid': False,
                    'error': f'File too large. Maximum size: {self.max_file_size / 1024 / 1024}MB'
                }

            # Check file extension
            if allowed_extensions:
                file_ext = Path(filename).suffix.lower()
                if file_ext not in allowed_extensions:
                    return {
                        'valid': False,
                        'error': f'File type not allowed. Allowed: {", ".join(allowed_extensions)}'
                    }

            # Detect MIME type
            mime_type = self._detect_mime_type(file)

            return {
                'valid': True,
                'filename': filename,
                'size': file_size,
                'mime_type': mime_type,
                'extension': Path(filename).suffix.lower()
            }

        except Exception as e:
            logger.error(f"❌ Error validating file: {e}")
            return {'valid': False, 'error': str(e)}

    def _detect_mime_type(self, file) -> str:
        """Detect MIME type using multiple methods"""
        try:
            # Method 1: Use python-magic if available
            try:
                import magic
                file.seek(0)
                mime_type = magic.from_buffer(file.read(1024), mime=True)
                file.seek(0)
                return mime_type
            except ImportError:
                pass

            # Method 2: Use mimetypes
            if hasattr(file, 'filename'):
                mime_type, _ = mimetypes.guess_type(file.filename)
                if mime_type:
                    return mime_type

            # Method 3: Default fallback
            ext = Path(file.filename).suffix.lower()
            mime_map = {
                '.pdf': 'application/pdf',
                '.doc': 'application/msword',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.txt': 'text/plain',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.gif': 'image/gif',
                '.mp4': 'video/mp4',
                '.avi': 'video/x-msvideo',
                '.mp3': 'audio/mpeg',
                '.wav': 'audio/wav'
            }

            return mime_map.get(ext, 'application/octet-stream')

        except Exception as e:
            logger.error(f"❌ Error detecting MIME type: {e}")
            return 'application/octet-stream'

    def upload_file(self, file, user_id: str = 'anonymous',
                   category: str = 'general',
                   metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """Upload and process a file"""
        try:
            # Validate file
            validation = self.validate_file(file)
            if not validation['valid']:
                return validation

            # Generate unique ID
            content_id = str(uuid.uuid4())

            # Create file path
            file_path = self._get_file_path(content_id, validation['filename'])

            # Ensure directory exists
            file_path.parent.mkdir(parents=True, exist_ok=True)

            # Save file
            file.seek(0)
            with open(file_path, 'wb') as f:
                shutil.copyfileobj(file, f)

            # Extract metadata
            extracted_metadata = self._extract_metadata(file_path, validation)

            # Create content record
            content_record = {
                'id': content_id,
                'filename': validation['filename'],
                'original_filename': file.filename,
                'path': str(file_path),
                'size': validation['size'],
                'mime_type': validation['mime_type'],
                'extension': validation['extension'],
                'category': category,
                'user_id': user_id,
                'uploaded_at': datetime.now().isoformat(),
                'hash': self._calculate_file_hash(file_path),
                'metadata': {
                    **extracted_metadata,
                    **(metadata or {})
                },
                'status': 'active',
                'version': 1,
                'download_count': 0,
                'last_accessed': datetime.now().isoformat()
            }

            # Store in database
            self.content_db[content_id] = content_record
            self._save_content_database()

            logger.info(f"✅ File uploaded successfully: {content_id} - {validation['filename']}")

            return {
                'success': True,
                'content_id': content_id,
                'content': content_record
            }

        except Exception as e:
            logger.error(f"❌ Error uploading file: {e}")
            return {'success': False, 'error': str(e)}

    def _get_file_path(self, content_id: str, filename: str) -> Path:
        """Generate file path for content"""
        # Create organized directory structure
        date_folder = datetime.now().strftime('%Y/%m/%d')
        file_path = self.upload_folder / date_folder / content_id

        # Add filename
        return file_path / filename

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, 'rb') as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"❌ Error calculating file hash: {e}")
            return ""

    def _extract_metadata(self, file_path: Path, file_info: Dict[str, Any]) -> Dict[str, Any]:
        """Extract metadata from file"""
        metadata = {
            'dimensions': None,
            'duration': None,
            'page_count': None,
            'word_count': None,
            'character_count': None,
            'encoding': None,
            'language': None
        }

        try:
            mime_type = file_info['mime_type']

            if mime_type.startswith('image/'):
                metadata.update(self._extract_image_metadata(file_path))

            elif mime_type == 'application/pdf':
                metadata.update(self._extract_pdf_metadata(file_path))

            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                              'application/msword']:
                metadata.update(self._extract_docx_metadata(file_path))

            elif mime_type.startswith('text/'):
                metadata.update(self._extract_text_metadata(file_path))

            elif mime_type.startswith('video/') or mime_type.startswith('audio/'):
                metadata.update(self._extract_media_metadata(file_path))

        except Exception as e:
            logger.error(f"❌ Error extracting metadata: {e}")

        return metadata

    def _extract_image_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from image files"""
        try:
            with Image.open(file_path) as img:
                return {
                    'dimensions': {
                        'width': img.width,
                        'height': img.height
                    },
                    'format': img.format,
                    'mode': img.mode,
                    'has_exif': hasattr(img, 'info') and img.info
                }
        except Exception as e:
            logger.error(f"❌ Error extracting image metadata: {e}")
            return {}

    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files"""
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                return {
                    'page_count': len(pdf_reader.pages),
                    'is_encrypted': pdf_reader.is_encrypted,
                    'metadata': pdf_reader.metadata.__dict__ if pdf_reader.metadata else {}
                }
        except Exception as e:
            logger.error(f"❌ Error extracting PDF metadata: {e}")
            return {}

    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

            return {
                'page_count': len(doc.sections),
                'paragraph_count': len(doc.paragraphs),
                'word_count': len(text.split()),
                'character_count': len(text)
            }
        except Exception as e:
            logger.error(f"❌ Error extracting DOCX metadata: {e}")
            return {}

    def _extract_text_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from text files"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

            lines = text.split('\n')
            words = text.split()
            chars = len(text)

            return {
                'line_count': len(lines),
                'word_count': len(words),
                'character_count': chars,
                'encoding': 'utf-8',
                'has_unicode': any(ord(c) > 127 for c in text)
            }
        except Exception as e:
            logger.error(f"❌ Error extracting text metadata: {e}")
            return {}

    def _extract_media_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from media files (video/audio)"""
        try:
            # For now, return basic file info
            # In production, you might use libraries like moviepy or mutagen
            return {
                'duration': None,  # Would need media library
                'bitrate': None,
                'codec': None,
                'sample_rate': None
            }
        except Exception as e:
            logger.error(f"❌ Error extracting media metadata: {e}")
            return {}

    def get_content(self, content_id: str) -> Optional[Dict[str, Any]]:
        """Get content by ID"""
        try:
            if content_id in self.content_db:
                content = self.content_db[content_id].copy()
                # Update last accessed
                content['last_accessed'] = datetime.now().isoformat()
                content['download_count'] += 1
                self.content_db[content_id] = content
                self._save_content_database()
                return content
            return None
        except Exception as e:
            logger.error(f"❌ Error getting content {content_id}: {e}")
            return None

    def list_content(self, user_id: str = None,
                    category: str = None,
                    limit: int = 50,
                    offset: int = 0) -> Dict[str, Any]:
        """List content with filtering"""
        try:
            content_list = list(self.content_db.values())

            # Apply filters
            if user_id:
                content_list = [c for c in content_list if c['user_id'] == user_id]

            if category:
                content_list = [c for c in content_list if c['category'] == category]

            # Sort by upload date (newest first)
            content_list.sort(key=lambda x: x['uploaded_at'], reverse=True)

            # Apply pagination
            total = len(content_list)
            start = offset
            end = offset + limit
            paginated_content = content_list[start:end]

            return {
                'content': paginated_content,
                'total': total,
                'limit': limit,
                'offset': offset,
                'has_more': end < total
            }

        except Exception as e:
            logger.error(f"❌ Error listing content: {e}")
            return {'content': [], 'total': 0, 'error': str(e)}

    def search_content(self, query: str,
                      user_id: str = None,
                      category: str = None) -> List[Dict[str, Any]]:
        """Search content by filename, metadata, or content"""
        try:
            query_lower = query.lower()
            results = []

            for content_id, content in self.content_db.items():
                # Check basic fields
                if (query_lower in content['filename'].lower() or
                    query_lower in content.get('original_filename', '').lower()):

                    # Apply filters
                    if user_id and content['user_id'] != user_id:
                        continue
                    if category and content['category'] != category:
                        continue

                    results.append(content)
                    continue

                # Check metadata
                metadata = content.get('metadata', {})
                if any(query_lower in str(value).lower()
                      for value in metadata.values() if value):
                    results.append(content)

            return results[:50]  # Limit results

        except Exception as e:
            logger.error(f"❌ Error searching content: {e}")
            return []

    def delete_content(self, content_id: str, user_id: str = None) -> Dict[str, Any]:
        """Delete content"""
        try:
            if content_id not in self.content_db:
                return {'success': False, 'error': 'Content not found'}

            content = self.content_db[content_id]

            # Check permissions
            if user_id and content['user_id'] != user_id:
                return {'success': False, 'error': 'Permission denied'}

            # Delete file
            file_path = Path(content['path'])
            if file_path.exists():
                file_path.unlink()

            # Remove from database
            del self.content_db[content_id]
            self._save_content_database()

            logger.info(f"✅ Content deleted: {content_id}")
            return {'success': True, 'content_id': content_id}

        except Exception as e:
            logger.error(f"❌ Error deleting content {content_id}: {e}")
            return {'success': False, 'error': str(e)}

    def get_content_file(self, content_id: str) -> Optional[Tuple[Path, str]]:
        """Get content file path and filename"""
        try:
            content = self.get_content(content_id)
            if content:
                file_path = Path(content['path'])
                if file_path.exists():
                    return file_path, content['filename']
            return None
        except Exception as e:
            logger.error(f"❌ Error getting content file {content_id}: {e}")
            return None

    def get_content_stream(self, content_id: str) -> Optional[BinaryIO]:
        """Get content file stream for serving"""
        try:
            result = self.get_content_file(content_id)
            if result:
                file_path, filename = result
                return open(file_path, 'rb')
            return None
        except Exception as e:
            logger.error(f"❌ Error getting content stream {content_id}: {e}")
            return None

    def get_content_stats(self) -> Dict[str, Any]:
        """Get content statistics"""
        try:
            total_content = len(self.content_db)
            total_size = sum(c['size'] for c in self.content_db.values())
            categories = {}
            file_types = {}

            for content in self.content_db.values():
                # Category stats
                category = content['category']
                categories[category] = categories.get(category, 0) + 1

                # File type stats
                ext = content['extension']
                file_types[ext] = file_types.get(ext, 0) + 1

            return {
                'total_content': total_content,
                'total_size': total_size,
                'total_size_mb': total_size / 1024 / 1024,
                'categories': categories,
                'file_types': file_types,
                'last_updated': datetime.now().isoformat()
            }

        except Exception as e:
            logger.error(f"❌ Error getting content stats: {e}")
            return {'error': str(e)}

    def cleanup_temp_files(self, older_than_hours: int = 24):
        """Clean up temporary files"""
        try:
            cutoff_time = datetime.now().timestamp() - (older_than_hours * 3600)

            cleaned = 0
            for file_path in self.temp_folder.glob("**/*"):
                if file_path.is_file() and file_path.stat().st_mtime < cutoff_time:
                    file_path.unlink()
                    cleaned += 1

            logger.info(f"🧹 Cleaned up {cleaned} temporary files older than {older_than_hours} hours")
            return {'success': True, 'cleaned_files': cleaned}

        except Exception as e:
            logger.error(f"❌ Error cleaning up temp files: {e}")
            return {'success': False, 'error': str(e)}

    def optimize_storage(self):
        """Optimize storage by removing unused files and organizing structure"""
        try:
            # Remove files not in database
            for file_path in self.upload_folder.glob("**/*"):
                if file_path.is_file() and file_path.name != "content_db.json":
                    # Check if file is referenced in database
                    file_referenced = False
                    for content in self.content_db.values():
                        if content['path'] == str(file_path):
                            file_referenced = True
                            break

                    if not file_referenced:
                        file_path.unlink()

            # Reorganize by date if needed
            # This is a basic implementation - in production you might want more sophisticated organization

            return {'success': True, 'message': 'Storage optimized'}

        except Exception as e:
            logger.error(f"❌ Error optimizing storage: {e}")
            return {'success': False, 'error': str(e)}

# Global instance
content_manager = ContentManager()

# Helper functions
def upload_file(file, user_id: str = 'anonymous', category: str = 'general', metadata: Dict[str, Any] = None) -> Dict[str, Any]:
    """Helper function to upload file"""
    return content_manager.upload_file(file, user_id, category, metadata)

def get_content(content_id: str) -> Optional[Dict[str, Any]]:
    """Helper function to get content"""
    return content_manager.get_content(content_id)

def list_content(user_id: str = None, category: str = None, limit: int = 50, offset: int = 0) -> Dict[str, Any]:
    """Helper function to list content"""
    return content_manager.list_content(user_id, category, limit, offset)

def search_content(query: str, user_id: str = None, category: str = None) -> List[Dict[str, Any]]:
    """Helper function to search content"""
    return content_manager.search_content(query, user_id, category)

def delete_content(content_id: str, user_id: str = None) -> Dict[str, Any]:
    """Helper function to delete content"""
    return content_manager.delete_content(content_id, user_id)

def get_content_file(content_id: str) -> Optional[Tuple[Path, str]]:
    """Helper function to get content file"""
    return content_manager.get_content_file(content_id)

def get_content_stats() -> Dict[str, Any]:
    """Helper function to get content statistics"""
    return content_manager.get_content_stats()

def cleanup_temp_files(older_than_hours: int = 24):
    """Helper function to cleanup temporary files"""
    return content_manager.cleanup_temp_files(older_than_hours)
