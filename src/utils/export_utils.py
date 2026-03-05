"""
Export Utilities for ValidoAI Application
Provides comprehensive export functionality for various formats
"""

import csv
import json
import xml.etree.ElementTree as ET
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import logging
import io
import base64
from pathlib import Path

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

logger = logging.getLogger(__name__)

class ExportUtils:
    """Utility class for data export operations"""
    
    @staticmethod
    def export_to_csv(data: List[Dict[str, Any]], 
                     filename: str = None,
                     delimiter: str = ',',
                     encoding: str = 'utf-8',
                     include_headers: bool = True) -> Union[str, bytes]:
        """Export data to CSV format"""
        try:
            if not data:
                return ""
            
            output = io.StringIO()
            fieldnames = list(data[0].keys()) if data else []
            
            writer = csv.DictWriter(
                output, 
                fieldnames=fieldnames,
                delimiter=delimiter
            )
            
            if include_headers:
                writer.writeheader()
            
            for row in data:
                writer.writerow(row)
            
            csv_content = output.getvalue()
            output.close()
            
            if filename:
                with open(filename, 'w', newline='', encoding=encoding) as f:
                    f.write(csv_content)
            
            return csv_content.encode(encoding)
        
        except Exception as e:
            logger.error(f"Error exporting to CSV: {e}")
            raise
    
    @staticmethod
    def export_to_excel(data: List[Dict[str, Any]], 
                       filename: str = None,
                       sheet_name: str = 'Sheet1',
                       include_headers: bool = True) -> Union[str, bytes]:
        """Export data to Excel format"""
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl is required for Excel export")
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            
            if not data:
                if filename:
                    wb.save(filename)
                return b""
            
            # Get headers
            headers = list(data[0].keys()) if data else []
            
            # Write headers
            if include_headers:
                for col, header in enumerate(headers, 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
            
            # Write data
            for row_idx, row_data in enumerate(data, 2):
                for col_idx, header in enumerate(headers, 1):
                    value = row_data.get(header, '')
                    ws.cell(row=row_idx, column=col_idx, value=value)
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            if filename:
                wb.save(filename)
                return filename
            
            # Return as bytes
            output = io.BytesIO()
            wb.save(output)
            output.seek(0)
            return output.getvalue()
        
        except Exception as e:
            logger.error(f"Error exporting to Excel: {e}")
            raise
    
    @staticmethod
    def export_to_word(data: List[Dict[str, Any]], 
                      filename: str = None,
                      title: str = "Data Export",
                      include_headers: bool = True) -> Union[str, bytes]:
        """Export data to Word format"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        try:
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_paragraph(title)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_paragraph.runs[0]
            title_run.font.size = 16
            title_run.font.bold = True
            
            doc.add_paragraph()  # Add spacing
            
            if not data:
                if filename:
                    doc.save(filename)
                return b""
            
            # Create table
            headers = list(data[0].keys()) if data else []
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            # Add headers
            if include_headers:
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = str(header)
                    header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Add data
            for row_data in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    value = row_data.get(header, '')
                    row_cells[i].text = str(value)
            
            if filename:
                doc.save(filename)
                return filename
            
            # Return as bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output.getvalue()
        
        except Exception as e:
            logger.error(f"Error exporting to Word: {e}")
            raise
    
    @staticmethod
    def export_to_pdf(data: List[Dict[str, Any]], 
                     filename: str = None,
                     title: str = "Data Export",
                     include_headers: bool = True) -> Union[str, bytes]:
        """Export data to PDF format"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export")
        
        try:
            if filename:
                doc = SimpleDocTemplate(filename, pagesize=A4)
            else:
                output = io.BytesIO()
                doc = SimpleDocTemplate(output, pagesize=A4)
            
            story = []
            styles = getSampleStyleSheet()
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            if not data:
                if filename:
                    doc.build(story)
                else:
                    doc.build(story)
                    output.seek(0)
                    return output.getvalue()
                return filename
            
            # Prepare table data
            headers = list(data[0].keys()) if data else []
            table_data = []
            
            # Add headers
            if include_headers:
                table_data.append([str(header) for header in headers])
            
            # Add data
            for row_data in data:
                table_data.append([str(row_data.get(header, '')) for header in headers])
            
            # Create table
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            
            if filename:
                doc.build(story)
                return filename
            else:
                doc.build(story)
                output.seek(0)
                return output.getvalue()
        
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    @staticmethod
    def export_to_json(data: List[Dict[str, Any]], 
                      filename: str = None,
                      indent: int = 2,
                      ensure_ascii: bool = False) -> Union[str, bytes]:
        """Export data to JSON format"""
        try:
            json_content = json.dumps(data, indent=indent, ensure_ascii=ensure_ascii, default=str)
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(json_content)
                return filename
            
            return json_content.encode('utf-8')
        
        except Exception as e:
            logger.error(f"Error exporting to JSON: {e}")
            raise
    
    @staticmethod
    def export_to_xml(data: List[Dict[str, Any]], 
                     filename: str = None,
                     root_name: str = 'data',
                     item_name: str = 'item') -> Union[str, bytes]:
        """Export data to XML format"""
        try:
            root = ET.Element(root_name)
            
            for item_data in data:
                item = ET.SubElement(root, item_name)
                for key, value in item_data.items():
                    # Create safe element name
                    safe_key = key.replace(' ', '_').replace('-', '_')
                    element = ET.SubElement(item, safe_key)
                    element.text = str(value)
            
            xml_content = ET.tostring(root, encoding='unicode', method='xml')
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(xml_content)
                return filename
            
            return xml_content.encode('utf-8')
        
        except Exception as e:
            logger.error(f"Error exporting to XML: {e}")
            raise
    
    @staticmethod
    def export_to_html(data: List[Dict[str, Any]], 
                      filename: str = None,
                      title: str = "Data Export",
                      include_headers: bool = True,
                      css_styles: str = None) -> Union[str, bytes]:
        """Export data to HTML format"""
        try:
            if not data:
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <title>{title}</title>
                    <meta charset="utf-8">
                    {f'<style>{css_styles}</style>' if css_styles else ''}
                </head>
                <body>
                    <h1>{title}</h1>
                    <p>No data available</p>
                </body>
                </html>
                """
            else:
                headers = list(data[0].keys()) if data else []
                
                # Default CSS styles
                default_css = """
                <style>
                    body { font-family: Arial, sans-serif; margin: 20px; }
                    table { border-collapse: collapse; width: 100%; margin-top: 20px; }
                    th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                    th { background-color: #f2f2f2; font-weight: bold; }
                    tr:nth-child(even) { background-color: #f9f9f9; }
                    tr:hover { background-color: #f5f5f5; }
                    h1 { color: #333; }
                </style>
                """
                
                css = css_styles if css_styles else default_css
                
                # Build table HTML
                table_html = "<table>\n"
                
                # Add headers
                if include_headers:
                    table_html += "  <thead>\n    <tr>\n"
                    for header in headers:
                        table_html += f"      <th>{header}</th>\n"
                    table_html += "    </tr>\n  </thead>\n"
                
                # Add data
                table_html += "  <tbody>\n"
                for row_data in data:
                    table_html += "    <tr>\n"
                    for header in headers:
                        value = row_data.get(header, '')
                        table_html += f"      <td>{value}</td>\n"
                    table_html += "    </tr>\n"
                table_html += "  </tbody>\n</table>"
                
                html_content = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <title>{title}</title>
                    <meta charset="utf-8">
                    {css}
                </head>
                <body>
                    <h1>{title}</h1>
                    {table_html}
                </body>
                </html>
                """
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return filename
            
            return html_content.encode('utf-8')
        
        except Exception as e:
            logger.error(f"Error exporting to HTML: {e}")
            raise
    
    @staticmethod
    def export_to_text(data: List[Dict[str, Any]], 
                      filename: str = None,
                      delimiter: str = '\t',
                      include_headers: bool = True) -> Union[str, bytes]:
        """Export data to plain text format"""
        try:
            if not data:
                return ""
            
            output = io.StringIO()
            fieldnames = list(data[0].keys()) if data else []
            
            # Write headers
            if include_headers:
                output.write(delimiter.join(fieldnames) + '\n')
            
            # Write data
            for row in data:
                row_values = [str(row.get(field, '')) for field in fieldnames]
                output.write(delimiter.join(row_values) + '\n')
            
            text_content = output.getvalue()
            output.close()
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                return filename
            
            return text_content.encode('utf-8')
        
        except Exception as e:
            logger.error(f"Error exporting to text: {e}")
            raise

class ExportManager:
    """Manager class for handling multiple export operations"""
    
    def __init__(self):
        self.export_history = []
        self.export_templates = {}
    
    def export_data(self, data: List[Dict[str, Any]], 
                   format_type: str,
                   filename: str = None,
                   **kwargs) -> Union[str, bytes]:
        """Export data in specified format"""
        try:
            if format_type.lower() == 'csv':
                result = ExportUtils.export_to_csv(data, filename, **kwargs)
            elif format_type.lower() in ['excel', 'xlsx']:
                result = ExportUtils.export_to_excel(data, filename, **kwargs)
            elif format_type.lower() in ['word', 'docx']:
                result = ExportUtils.export_to_word(data, filename, **kwargs)
            elif format_type.lower() == 'pdf':
                result = ExportUtils.export_to_pdf(data, filename, **kwargs)
            elif format_type.lower() == 'json':
                result = ExportUtils.export_to_json(data, filename, **kwargs)
            elif format_type.lower() == 'xml':
                result = ExportUtils.export_to_xml(data, filename, **kwargs)
            elif format_type.lower() == 'html':
                result = ExportUtils.export_to_html(data, filename, **kwargs)
            elif format_type.lower() == 'text':
                result = ExportUtils.export_to_text(data, filename, **kwargs)
            else:
                raise ValueError(f"Unsupported export format: {format_type}")
            
            # Record export history
            self.export_history.append({
                'timestamp': datetime.now(),
                'format': format_type,
                'filename': filename,
                'record_count': len(data),
                'status': 'success'
            })
            
            return result
        
        except Exception as e:
            # Record failed export
            self.export_history.append({
                'timestamp': datetime.now(),
                'format': format_type,
                'filename': filename,
                'record_count': len(data),
                'status': 'failed',
                'error': str(e)
            })
            raise
    
    def batch_export(self, data: List[Dict[str, Any]], 
                    formats: List[str],
                    base_filename: str = None,
                    **kwargs) -> Dict[str, Union[str, bytes]]:
        """Export data in multiple formats"""
        results = {}
        
        for format_type in formats:
            try:
                filename = f"{base_filename}.{format_type}" if base_filename else None
                result = self.export_data(data, format_type, filename, **kwargs)
                results[format_type] = result
            except Exception as e:
                logger.error(f"Failed to export {format_type}: {e}")
                results[format_type] = None
        
        return results
    
    def get_export_history(self, limit: int = 100) -> List[Dict[str, Any]]:
        """Get export history"""
        return self.export_history[-limit:] if limit else self.export_history
    
    def clear_export_history(self):
        """Clear export history"""
        self.export_history.clear()

# Global export manager instance
export_manager = ExportManager()

# Convenience functions
def export_to_csv(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to CSV format"""
    return export_manager.export_data(data, 'csv', **kwargs)

def export_to_excel(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to Excel format"""
    return export_manager.export_data(data, 'excel', **kwargs)

def export_to_word(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to Word format"""
    return export_manager.export_data(data, 'word', **kwargs)

def export_to_pdf(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to PDF format"""
    return export_manager.export_data(data, 'pdf', **kwargs)

def export_to_json(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to JSON format"""
    return export_manager.export_data(data, 'json', **kwargs)

def export_to_xml(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to XML format"""
    return export_manager.export_data(data, 'xml', **kwargs)

def export_to_html(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to HTML format"""
    return export_manager.export_data(data, 'html', **kwargs)

def export_to_text(data: List[Dict[str, Any]], **kwargs) -> Union[str, bytes]:
    """Export data to text format"""
    return export_manager.export_data(data, 'text', **kwargs)
